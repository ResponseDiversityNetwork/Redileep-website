from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "share_packet/ReDiLEEP_website_creation_options_brief.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(22, 36, 33)
MUTED = RGBColor(90, 105, 100)
LIGHT_GRAY = "F2F4F7"
LIGHT_TEAL = "E9F1EA"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_tbl_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_border_bottom(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_hyperlink_like(doc, label, target):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    set_run_font(r1, bold=True)
    r2 = p.add_run(target)
    set_run_font(r2, color=BLUE)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_tbl_width(table, [6.35])
    set_table_borders(table, color="BFD6CE", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=INK)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, color=INK)


def add_comparison_table(doc):
    rows = [
        ("Best fit", "Static, content-led research network site with controlled design, fast publishing, low technical overhead once workflow is set.", "Larger interactive site with many non-technical editors, frequent ad-hoc pages, plugins, forms, and built-in CMS workflows."),
        ("Set-up", "AI can scaffold a polished site quickly; developer/maintainer then reviews, refines, and publishes through GitHub/Quarto hosting.", "Theme selection, hosting, plugin stack, page builder configuration, security settings, and editorial permissions."),
        ("Maintenance", "Edit text in Markdown/QMD files; AI can make updates from plain-language requests; no server patching if hosted statically.", "Editors use browser UI; site owner must maintain WordPress core, theme, plugins, backups, updates, and security."),
        ("Multi-person contribution", "Works well with GitHub pull requests, issue templates, or a light editorial intake form. Can support several contributors with review before publish.", "Very familiar CMS roles for many editors. Easier for direct browser editing, but quality control can drift without governance."),
        ("Design control", "High control and reproducibility. Changes can be reviewed before publication and kept consistent.", "Depends heavily on theme/page-builder discipline. Easy to make local edits that gradually reduce consistency."),
        ("Security", "Static site has a small attack surface and no login database on the public site.", "Needs active security management, because login pages, plugins, and databases are common targets."),
        ("Cost", "Likely low hosting cost; main cost is initial build, governance, and occasional technical support.", "Hosting may be modest, but plugin subscriptions, updates, and maintenance can add recurring cost."),
        ("Risks", "Requires agreed workflow for non-coders; some changes may need AI/developer assistance; Git can feel unfamiliar.", "Plugin dependency, update breakage, security exposure, and long-term page-builder lock-in."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(["Question", "Code/AI Quarto site", "WordPress"]):
        set_cell_shading(hdr[i], LIGHT_GRAY)
        p = hdr[i].paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(text)
            set_run_font(r, size=9.5)
    set_tbl_width(table, [1.15, 2.6, 2.6])


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_borders(table, color="BFD6CE", size="6")
    hdr = table.rows[0].cells
    for i, text in enumerate(["Item", "Draft outcome"]):
        set_cell_shading(hdr[i], LIGHT_TEAL)
        r = hdr[i].paragraphs[0].add_run(text)
        set_run_font(r, bold=True)
    rows = [
        ("Prompt", "A minimal request: make a Quarto website for a newly funded MSCA Doctoral Network."),
        ("Elapsed time", "A working draft was created in one coding session, roughly under an hour from empty folder to rendered, browser-checked site."),
        ("Inputs used", "Project folder name, the minimal prompt, then proposal PDFs for public-facing ReDiLEEP details."),
        ("Outputs", "Quarto source files, rendered static website, generated hero image, navigation, draft pages, news item, and shareable static ZIP."),
        ("Human review still needed", "Official wording, partner approvals, recruitment details, accessibility/legal/privacy checks, hosting choice, maintenance workflow."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            r = cells[i].paragraphs[0].add_run(text)
            set_run_font(r, size=10)
    set_tbl_width(table, [1.45, 4.9])


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Title", 23, INK, 0, 4),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "ReDiLEEP website options brief"
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = MUTED

footer = section.footer.paragraphs[0]
footer.text = "Draft for partner discussion"
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = MUTED

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Decision brief")
set_run_font(r, size=11, bold=True, color=MUTED)

p = doc.add_paragraph(style="Title")
r = p.add_run("ReDiLEEP Draft Website: Creation Approach, Maintenance Options, and WordPress Comparison")
set_run_font(r, size=23, bold=True, color=INK)

p = doc.add_paragraph()
r = p.add_run("Prepared for ReDiLEEP project partners and leaders | Draft for discussion | 21 May 2026")
set_run_font(r, size=10.5, color=MUTED)
paragraph_border_bottom(p, color="0D766E", size="12")

doc.add_heading("Executive Summary", level=1)
add_callout(
    doc,
    "Short recommendation",
    "The code/AI-generated Quarto approach is credible for a polished, low-maintenance MSCA network website, provided the consortium first agrees a light design and operational requirements specification. WordPress remains attractive if the priority is browser-based editing by many non-technical contributors without a reviewed publishing workflow."
)
add_bullets(doc, [
    "The current draft was created from a very small initial prompt and then refined using project-specific information from the ReDiLEEP proposal PDFs.",
    "The draft demonstrates that a static Quarto website can be generated quickly, styled coherently, and packaged so partners can view it without installing software.",
    "The main decision is not only technology; it is governance: who can edit, who approves, how recruitment/news updates are handled, and how long-term hosting/security are managed.",
    "A code/AI approach can be maintainer-friendly if updates are channelled through simple Markdown files, issue/request forms, shared content templates, and AI-assisted edits reviewed before publication.",
])

doc.add_heading("Contents", level=1)
add_numbered(doc, [
    "Suggested covering note",
    "What was created and how to view it",
    "How the draft was created",
    "What next if we proceed with a code/AI website",
    "WordPress versus code/AI-generated Quarto",
    "Maintainer and contributor friendliness",
    "Suggested decision points",
])

doc.add_heading("1. Suggested Covering Note", level=1)
add_para(doc, "Dear colleagues,")
add_para(doc, "I am sharing a very early draft ReDiLEEP website for discussion, together with this short note on how it was created and what the options might be for future maintenance. The draft was produced from a deliberately minimal prompt asking for a Quarto website for a newly funded MSCA Doctoral Network, then refined using high-level project information. The initial working draft was created in one coding session, roughly under an hour from empty folder to rendered, browser-checked static site.")
add_para(doc, "The purpose is not to suggest that this draft is ready for publication. Rather, it is a quick demonstration of what a code/AI-assisted website workflow could produce, and a basis for comparing that route with WordPress. If we wanted to explore this seriously, the next step would be to agree a design and operational requirements specification together before deciding on the platform.")
add_para(doc, "To view the draft site, unzip the website package and open index.html. No software installation is needed.")

doc.add_heading("2. What Was Created And How To View It", level=1)
add_summary_table(doc)
add_hyperlink_like(doc, "Static website package", "redileep-draft-website-static.zip")
add_para(doc, "After unzipping, partners can open index.html directly in a browser. This is intended for review only; it is not yet a production deployment.")

doc.add_heading("3. How The Draft Was Created", level=1)
add_para(doc, "The initial instruction was deliberately minimal: “please can you make a quarto website for a mcsa doctoral network that we just got funded”. From an empty Git folder, the workflow was:")
add_numbered(doc, [
    "Create a Quarto website structure, navigation, pages, styling, favicon, and static output configuration.",
    "Generate a hero image suitable for a science/research network website and copy it into the project assets.",
    "Render the site with Quarto and inspect it in a browser preview.",
    "Use the proposal PDFs to replace generic placeholders with high-level ReDiLEEP-specific information: response diversity focus, acronym expansion, work packages, partners, and 13 doctoral candidate projects.",
    "Re-render and browser-check the site, then package the static output for sharing.",
])
add_para(doc, "The important point for the consortium is that the AI did not need a long specification to create a coherent first draft. However, a production-quality website should not rely on improvisation. The next phase should be guided by an agreed specification.")

doc.add_heading("4. Next Steps If We Proceed With A Code/AI Website", level=1)
add_numbered(doc, [
    "Create a design and operational requirements specification with partners: audience, tone, visual references, information architecture, languages, accessibility, approvals, update frequency, and legal requirements.",
    "Define the editorial model: who drafts news, who updates vacancies, who approves partner profiles, and who publishes.",
    "Convert proposal-derived draft text into approved public copy, including EU funding acknowledgement, grant agreement number, recruitment wording, partner descriptions, and contact details.",
    "Choose hosting and deployment: for example GitHub Pages, institutional hosting, Netlify, or another static host.",
    "Set up a contribution workflow: GitHub issues or forms for requests, protected main branch, preview links for proposed changes, and a simple release checklist.",
    "Add production essentials: accessibility pass, privacy/cookie/legal pages, analytics decision, search behaviour, metadata/social previews, image alt text, and backup/domain ownership plan.",
    "Train one or two maintainers and document routine updates such as adding a news post, changing a vacancy, or editing partner details.",
])

doc.add_heading("5. WordPress Versus Code/AI-Generated Quarto", level=1)
add_comparison_table(doc)

doc.add_heading("6. Maintainer And Contributor Friendliness", level=1)
doc.add_heading("How friendly can a code/AI site be?", level=2)
add_para(doc, "Very friendly, if the workflow is designed around the actual maintainers rather than around software ideology. Most content can live in plain-text Markdown/QMD files with simple front matter. A maintainer can ask an AI assistant to “add this news item”, “update the vacancies page”, or “make a partner profile from this text”, then review the exact change before it is published.")
add_bullets(doc, [
    "Routine content can be template-based: news posts, vacancy pages, partner cards, events, deliverables, and announcements.",
    "Changes can be reviewed through previews before publication, reducing accidental layout or wording drift.",
    "The public site can remain static, fast, and secure even if editing happens through GitHub, a desktop editor, or AI-assisted workflows.",
    "For non-coders, the workflow should hide Git complexity as much as possible: request forms, checklists, named maintainers, and clear review rules.",
])

doc.add_heading("Multiple contributors", level=2)
add_para(doc, "A code/AI site can support multiple contributors, but it should use a reviewed contribution model rather than unrestricted direct editing. Suitable models include:")
add_bullets(doc, [
    "Central maintainer model: partners email/request changes; one or two maintainers update and publish.",
    "Contributor model: approved contributors edit Markdown files or submit content through forms; maintainers approve and merge.",
    "Hybrid model: content owners draft in Word/Google Docs, AI converts to site format, maintainers review layout and publish.",
])
add_para(doc, "WordPress is more immediately familiar for many editors because it offers browser-based editing and roles. The trade-off is that more direct editors often means more training, more design drift, and more need for content governance.")

doc.add_heading("7. Suggested Decision Points", level=1)
add_bullets(doc, [
    "Do we want a small, polished, largely static public website, or a full CMS with many direct editors?",
    "How many people need to publish directly, rather than submit content for review?",
    "Who will own security, updates, backups, and continuity after the project launches?",
    "How important are low hosting/security overhead and reproducible design?",
    "Are partners comfortable with an AI-assisted maintenance workflow if routine updates are made through simple requests and reviewed previews?",
])

doc.add_heading("Proposed Immediate Action", level=1)
add_callout(
    doc,
    "One practical next step",
    "Use the shared draft website and this brief for a short partner discussion. If there is interest, hold a 60-90 minute requirements workshop and produce a design and operational requirements specification before deciding finally between WordPress and a code/AI-generated Quarto site."
)

doc.save(OUT)
print(OUT)
